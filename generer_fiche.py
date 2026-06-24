import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def completer_fiche_word(texte_ia, nom_fichier="Fiche_Chantier_Wagner.docx"):
    doc = Document()
    
    # Configuration des marges et du style par défaut
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x34, 0x49, 0x5E) # Gris foncé pro

    # Titre principal du document
    titre = doc.add_paragraph()
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titre = titre.add_run("FICHE DE QUALIFICATION DE CHANTIER")
    run_titre.bold = True
    run_titre.font.size = Pt(18)
    run_titre.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    
    doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph() # Espace vide

    # Découpage du texte de l'IA ligne par ligne
    lignes = texte_ia.split('\n')
    
    for ligne in lignes:
        ligne_propre = ligne.replace("**", "").replace("*", "").strip()
        
        if not ligne_propre:
            continue
            
        # Si c'est une section principale (ex: 1. FICHE CLIENT, 4. DESCRIPTIF...)
        if any(sec in ligne_propre for sec in ["1. FICHE CLIENT", "2. INFOS RDV", "3. CONTEXTE IMMOBILIER", "4. DESCRIPTIF", "5. POINTS DE VIGILANCE"]):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(ligne_propre)
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x29, 0x80, 0xB9) # Bleu pro
            
        # Si c'est un sous-lot (ex: LOT Électricité, LOT Plâtrerie)
        elif "LOT " in ligne_propre:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            run = p.add_run(ligne_propre)
            run.bold = True
            run.font.color.rgb = RGBColor(0x16, 0xA0, 0x85) # Vert émeraude technique
            
        # Si c'est une ligne de détails ou une puce
        else:
            # On détecte si c'est une sous-puce ou une info de type "Clé : Valeur"
            if ":" in ligne_propre:
                p = doc.add_paragraph(style='Normal')
                cle, valeur = ligne_propre.split(":", 1)
                run_cle = p.add_run(cle + ":")
                run_cle.bold = True
                p.add_run(valeur)
            else:
                doc.add_paragraph(ligne_propre, style='List Bullet')

    # Sauvegarde du fichier
    doc.save(nom_fichier)
    print(f"🎉 Le fichier Word a été généré avec succès : '{nom_fichier}'")